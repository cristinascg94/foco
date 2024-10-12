from django.contrib import admin
from .models import  Corto, Pase, UsuarioAleatorio, Votacion
from docx import Document
import io
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.http import HttpResponse
import qrcode
from io import BytesIO
from docx.shared import Inches
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def crear_codigo_qr(url):
    """
    Crea un código QR para una URL y devuelve la imagen como un objeto BytesIO.

    :param url: La URL que se va a codificar en el código QR.
    :return: Un objeto BytesIO que contiene la imagen del código QR.
    """
    # Crear el objeto QR
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )

    # Agregar la URL al código QR
    qr.add_data(url)
    qr.make(fit=True)

    # Crear la imagen del código QR
    img = qr.make_image(fill_color="black", back_color="white")

    # Guardar la imagen en un objeto BytesIO
    img_bytes = BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)  # Mover el puntero al inicio del objeto BytesIO

    return img_bytes


def set_cell_border(cell, **kwargs):
    """
    Función para establecer bordes en las celdas.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Los bordes posibles son 'top', 'left', 'bottom', 'right', 'insideH', 'insideV'
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_elem = OxmlElement(f'w:{edge}')
        edge_elem.set(qn('w:val'), 'single')
        edge_elem.set(qn('w:sz'), '8')  # Tamaño del borde (puedes ajustarlo)
        edge_elem.set(qn('w:space'), '0')
        edge_elem.set(qn('w:color'), '000000')  # Color del borde en hexadecimal
        tcPr.append(edge_elem)

def generar_word(modeladmin, request, queryset):
    # Filtrar los usuarios
    for obj in queryset:
        print(obj.pase)
        pase = Pase.objects.get(pase=obj.pase)
        usuarios_filtrados = UsuarioAleatorio.objects.filter(pase=pase)

        # Crear el documento Word
        doc = Document()

        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0)    # Márgenes superior
            section.bottom_margin = Inches(0) # Márgenes inferior
            section.left_margin = Inches(0)   # Márgenes izquierdo
            section.right_margin = Inches(0)  # Márgenes derecho


        # Agregar tabla con dos columnas
        table = doc.add_table(rows=1, cols=4)

        n=0
        qr_url = f"http://13.60.206.34/votacion/" 
        qr_image = crear_codigo_qr(qr_url)

        usuarios_escribir = []
        for usuario in usuarios_filtrados:

            usuarios_escribir.append(usuario)
            
            if len(usuarios_escribir)==4:
                row_cells = table.add_row().cells
                for n_user, user in enumerate(usuarios_escribir):
                    row_cells[n_user].text = f"\n\nPase: {pase}\nUsuario: {user}"
                    run = row_cells[n_user].add_paragraph().add_run()  # Crear un nuevo párrafo en la celda
                    run.add_picture(qr_image, width=Inches(1.0))
                    for paragraph in row_cells[n_user].paragraphs:
                        paragraph.alignment = 1  # 1 es para centrar

                usuarios_escribir = []

                # Aplicar bordes a cada celda
                for cell in row_cells:
                    set_cell_border(cell)
            n_hoja =n+1
            if n_hoja%16 == 0:
                print(f"entro {n_hoja}")
                doc.add_page_break()
                table = doc.add_table(rows=1, cols=4)
            print(n)
            n=n+1

        # Preparar la respuesta HTTP
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)

        response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=usuarios_con_pase_{pase}.docx'

    return response

generar_word.short_description = "generar_word"


admin.site.register(Corto)

admin.site.register(UsuarioAleatorio)

class PaseAdmin(admin.ModelAdmin):
    actions = [generar_word]
    list_display = ('pase',)

admin.site.register(Pase, PaseAdmin)


class VotacionAdmin(admin.ModelAdmin):
    # Muestra las columnas 'corto', 'usuario' y 'votacion' en la lista del admin
    list_display = ('corto', 'usuario', 'votacion')
admin.site.register(Votacion, VotacionAdmin)